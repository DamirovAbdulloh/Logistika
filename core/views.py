from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.db.models import Q, Sum
import re
import os
from django.http import HttpResponse
from django.utils import timezone
from datetime import date, timedelta
from decimal import Decimal
from functools import wraps

from .models import Driver, Putyovka, TIR, Dazvol, Litsenziya, IjaraShartnoma, Chiqim, TolovQayd
from .forms import (AdminLoginForm, DriverLoginForm, DriverForm,
                    PutyovkaForm, TIRForm, DazvolForm, LitsenziyaForm,
                    IjaraForm, ChiqimForm, TolovForm)

def create_initial_payment(obj, tur):
    if getattr(obj,'tolangan',0) and obj.tolangan>0:
        TolovQayd.objects.create(driver=obj.driver,tur=tur,summa=obj.tolangan,izoh="Boshlang'ich to'lov")


# ─── Decorators ───────────────────────────────────────────────────────────────

def admin_required(view_func):
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated or not request.user.is_staff:
            return redirect('login')
        return view_func(request, *args, **kwargs)
    return wrapper


def driver_required(view_func):
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.session.get('driver_id'):
            return redirect('login')
        return view_func(request, *args, **kwargs)
    return wrapper


# ─── Auth Views ───────────────────────────────────────────────────────────────


def login_view(request):
    if request.user.is_authenticated:
        return redirect('admin_dashboard')

    if request.session.get('driver_id'):
        return redirect('driver_dashboard')

    admin_form = AdminLoginForm()
    driver_form = DriverLoginForm()

    if request.method == 'POST':
        role = request.POST.get('role')

        # ===== ADMIN LOGIN =====
        if role == 'admin':
            admin_form = AdminLoginForm(request.POST)

            if admin_form.is_valid():
                user = authenticate(
                    username=admin_form.cleaned_data['username'],
                    password=admin_form.cleaned_data['password']
                )

                if user and user.is_staff:
                    login(request, user)
                    return redirect('admin_dashboard')

                messages.error(request, "Login yoki parol noto'g'ri!")

        # ===== DRIVER LOGIN =====
        elif role == 'driver':
            driver_form = DriverLoginForm(request.POST)

            if driver_form.is_valid():
                telefon = driver_form.cleaned_data['telefon'].strip()
                parol = driver_form.cleaned_data['parol'].strip()

                # +998 90 123 45 67 -> +998901234567
                telefon = re.sub(r'\D', '', telefon)

                if len(telefon) == 9:
                    telefon = '998' + telefon

                if telefon.startswith('998'):
                    telefon = '+' + telefon

                driver = Driver.objects.filter(
                    telefon=telefon,
                    parol=parol
                ).first()

                if driver:
                    request.session['driver_id'] = driver.id
                    request.session['driver_ism'] = driver.ism

                    return redirect('driver_dashboard')

                messages.error(request, "Telefon yoki parol noto'g'ri!")

    return render(request, 'auth/login.html', {
        'admin_form': admin_form,
        'driver_form': driver_form,
    })
    
    
def logout_view(request):
    logout(request)
    request.session.flush()
    return redirect('login')


# ─── Admin Views ──────────────────────────────────────────────────────────────

@admin_required
def admin_dashboard(request):
    today = date.today()
    soon = today + timedelta(days=30)
    month_start = today.replace(day=1)

    # Stats
    stats = {
        'haydovchilar':    Driver.objects.count(),
        'adminlar':        User.objects.filter(is_staff=True).count(),
        'profillari':      Driver.objects.exclude(mashina='').count(),
        'putyovkalar':     Putyovka.objects.count(),
        'putyovka_tugagan': Putyovka.objects.filter(status='tugagan').count(),
        'tirlar':          TIR.objects.count(),
        'tir_tugagan':     TIR.objects.filter(tugash__lt=today).count(),
        'dazvollar':       Dazvol.objects.count(),
        'dazvol_tugagan':  Dazvol.objects.filter(tugash__lt=today).count(),
        'litsenziyalar':   Litsenziya.objects.count(),
        'ijara':           IjaraShartnoma.objects.filter(status='faol').count(),
    }

    # Kirim (income) - TolovQayd asosida
    kirim = {
        'kunlik': TolovQayd.objects.filter(sana=today).aggregate(s=Sum('summa'))['s'] or 0,
        'oylik':  TolovQayd.objects.filter(sana__gte=month_start).aggregate(s=Sum('summa'))['s'] or 0,
        'jami':   TolovQayd.objects.aggregate(s=Sum('summa'))['s'] or 0,
    }
    # Kunlik kirimlar - oxirgi 7 kun
    kunlik_kirimlar = []
    for i in range(6, -1, -1):
        d_ = today - timedelta(days=i)
        s = TolovQayd.objects.filter(sana=d_).aggregate(s=Sum('summa'))['s'] or 0
        kunlik_kirimlar.append({'sana': d_.strftime('%d.%m'), 'summa': s})

    # Oxirgi to'lovlar
    oxirgi_kirimlar = TolovQayd.objects.select_related('driver').order_by('-sana', '-id')[:8]

    # Qarzdor + faol haydovchilar (dashboardda haydovchi bo'limi)
    haydovchi_stats = []
    for driver in Driver.objects.all():
        haydovchi_stats.append({
            'id': driver.id,
            'ism': driver.ism,
            'mashina': driver.mashina,
            'qarz': driver.jami_qarz,
            'tolangan': driver.jami_tolangan,
            'aktiv': driver.aktiv_hujjatlar_soni,
        })
    # Eng ko'p qarzli — top 8
    qarzdorlar = sorted([h for h in haydovchi_stats if h['qarz'] > 0],
                        key=lambda x: x['qarz'], reverse=True)[:8]
    # Haydovchilar jadvali — top 10
    haydovchi_stats.sort(key=lambda x: x['aktiv'], reverse=True)
    haydovchi_top = haydovchi_stats[:10]

    # Muddati yaqin hujjatlar
    muddati_yaqin = []
    for p in Putyovka.objects.filter(tugash__gte=today, tugash__lte=soon, status='faol'):
        muddati_yaqin.append({'ism': p.driver.ism, 'tur': 'Putyovka', 'sana': p.tugash.strftime('%d.%m.%Y')})
    for t in TIR.objects.filter(tugash__gte=today, tugash__lte=soon).exclude(status='tugagan'):
        muddati_yaqin.append({'ism': t.driver.ism, 'tur': 'TIR', 'sana': t.tugash.strftime('%d.%m.%Y')})
    for d in Dazvol.objects.filter(tugash__gte=today, tugash__lte=soon).exclude(status='tugagan'):
        muddati_yaqin.append({'ism': d.driver.ism, 'tur': 'Dazvol', 'sana': d.tugash.strftime('%d.%m.%Y')})
    muddati_yaqin.sort(key=lambda x: x['sana'])
    muddati_yaqin = muddati_yaqin[:10]

    # Muddati allaqachon tugagan hujjatlar (dashboard alert uchun)
    expired_docs = []
    for p in Putyovka.objects.filter(tugash__lt=today, status='faol').select_related('driver')[:20]:
        expired_docs.append({'ism': p.driver.ism, 'tur': 'Putyovka', 'driver_id': p.driver.id})
    for t in TIR.objects.filter(tugash__lt=today).exclude(status='tugagan').select_related('driver')[:10]:
        expired_docs.append({'ism': t.driver.ism, 'tur': 'TIR', 'driver_id': t.driver.id})
    for d in Dazvol.objects.filter(tugash__lt=today).exclude(status='tugagan').select_related('driver')[:10]:
        expired_docs.append({'ism': d.driver.ism, 'tur': 'Dazvol', 'driver_id': d.driver.id})

    expired_count = len(expired_docs)

    from django.urls import reverse
    stat_cards = [
        {'icon': '👥', 'bg': '#dbeafe', 'label': 'Haydovchilar',  'value': stats['haydovchilar'],  'expired': None,                              'url': reverse('admin_drivers')},
        {'icon': '👤', 'bg': '#fce7f3', 'label': 'Adminlar',      'value': stats['adminlar'],      'expired': None,                              'url': reverse('admin_add_admin')},
        {'icon': '📄', 'bg': '#fef9c3', 'label': 'Putyovkalar',   'value': stats['putyovkalar'],   'expired': stats['putyovka_tugagan'] or None, 'url': reverse('admin_putyovkalar')},
        {'icon': '🚛', 'bg': '#ede9fe', 'label': 'TIRlar',        'value': stats['tirlar'],        'expired': stats['tir_tugagan'] or None,      'url': reverse('admin_tirlar')},
        {'icon': '📝', 'bg': '#fee2e2', 'label': 'Dazvollar',     'value': stats['dazvollar'],     'expired': stats['dazvol_tugagan'] or None,   'url': reverse('admin_dazvollar')},
        {'icon': '🛡', 'bg': '#cffafe', 'label': 'Litsenziyalar', 'value': stats['litsenziyalar'], 'expired': None,                              'url': reverse('admin_litsenziyalar')},
        {'icon': '🏠', 'bg': '#dcfce7', 'label': 'Ijara (Rent)',  'value': stats['ijara'],         'expired': None,                              'url': reverse('admin_ijara')},
        {'icon': '👤', 'bg': '#ccfbf1', 'label': 'Profillari',    'value': stats['profillari'],    'expired': None,                              'url': reverse('admin_drivers')},
    ]

    return render(request, 'admin/dashboard.html', {
        'stat_cards': stat_cards,
        'qarzdorlar': qarzdorlar,
        'muddati_yaqin': muddati_yaqin,
        'expired_docs': expired_docs,
        'expired_count': expired_count,
        'kirim': kirim,
        'kunlik_kirimlar': kunlik_kirimlar,
        'oxirgi_kirimlar': oxirgi_kirimlar,
        'haydovchi_top': haydovchi_top,
    })


# ─── Tugatish / Faollashtirish (umumiy) ───────────────────────────────────────

def _toggle_status(model, pk, new_status):
    obj = get_object_or_404(model, pk=pk)
    obj.status = new_status
    obj.save()
    return obj


@admin_required
def admin_putyovka_activate(request, pk):
    _toggle_status(Putyovka, pk, 'faol')
    messages.success(request, "Putyovka qayta faollashtirildi!")
    return redirect('admin_putyovkalar')


@admin_required
def admin_tir_finish(request, pk):
    _toggle_status(TIR, pk, 'tugagan')
    messages.success(request, "TIR tugatildi!")
    return redirect('admin_tirlar')


@admin_required
def admin_tir_activate(request, pk):
    _toggle_status(TIR, pk, 'faol')
    messages.success(request, "TIR qayta faollashtirildi!")
    return redirect('admin_tirlar')


@admin_required
def admin_dazvol_finish(request, pk):
    _toggle_status(Dazvol, pk, 'tugagan')
    messages.success(request, "Dazvol tugatildi!")
    return redirect('admin_dazvollar')


@admin_required
def admin_dazvol_activate(request, pk):
    _toggle_status(Dazvol, pk, 'faol')
    messages.success(request, "Dazvol qayta faollashtirildi!")
    return redirect('admin_dazvollar')


@admin_required
def admin_litsenziya_finish(request, pk):
    _toggle_status(Litsenziya, pk, 'tugagan')
    messages.success(request, "Litsenziya tugatildi!")
    return redirect('admin_litsenziyalar')


@admin_required
def admin_litsenziya_activate(request, pk):
    _toggle_status(Litsenziya, pk, 'faol')
    messages.success(request, "Litsenziya qayta faollashtirildi!")
    return redirect('admin_litsenziyalar')


@admin_required
def admin_ijara_finish(request, pk):
    _toggle_status(IjaraShartnoma, pk, 'tugagan')
    messages.success(request, "Ijara tugatildi!")
    return redirect('admin_ijara')


@admin_required
def admin_ijara_activate(request, pk):
    _toggle_status(IjaraShartnoma, pk, 'faol')
    messages.success(request, "Ijara qayta faollashtirildi!")
    return redirect('admin_ijara')


@admin_required
def admin_drivers(request):
    today = date.today()
    q = request.GET.get('q', '')
    drivers = Driver.objects.all()
    if q:
        drivers = drivers.filter(
            Q(ism__icontains=q) | Q(telefon__icontains=q) |
            Q(mashina__icontains=q) | Q(raqam__icontains=q)
        )

    # Summary
    jami_qarz = sum(d.jami_qarz for d in Driver.objects.all())
    # TolovQayd dan hisoblaymiz — dashboard bilan bir xil manba
    jami_tolangan = TolovQayd.objects.aggregate(s=Sum('summa'))['s'] or 0
    jami_hujjat = (Putyovka.objects.count() + TIR.objects.count() +
                   Dazvol.objects.count() + Litsenziya.objects.count())
    aktiv = (Putyovka.objects.filter(status='faol').count() +
             TIR.objects.filter(tugash__gte=today).count())

    paginator = Paginator(drivers, 20)
    page_obj = paginator.get_page(request.GET.get('page'))

    return render(request, 'admin/drivers.html', {
        'drivers': page_obj,
        'search_q': q,
        'summary': {
            'jami': jami_hujjat,
            'tolangan': jami_tolangan,
            'qarz': jami_qarz,
            'aktiv': aktiv,
        }
    })


@admin_required
def admin_driver_add(request):
    if request.method == 'POST':
        form = DriverForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Haydovchi muvaffaqiyatli qo'shildi!")
            return redirect('admin_drivers')
    else:
        form = DriverForm()
    return render(request, 'admin/driver_form.html', {'form': form, 'title': "Haydovchi qo'shish"})


@admin_required
def admin_driver_edit(request, pk):
    driver = get_object_or_404(Driver, pk=pk)
    if request.method == 'POST':
        form = DriverForm(request.POST, instance=driver)
        if form.is_valid():
            form.save()
            messages.success(request, "Haydovchi ma'lumotlari yangilandi!")
            return redirect('admin_drivers')
    else:
        form = DriverForm(instance=driver)
    return render(request, 'admin/driver_form.html', {'form': form, 'title': "Haydovchini tahrirlash", 'driver': driver})


@admin_required
def admin_driver_delete(request, pk):
    driver = get_object_or_404(Driver, pk=pk)
    driver.delete()
    messages.success(request, "Haydovchi o'chirildi!")
    return redirect('admin_drivers')


@admin_required
def admin_driver_detail(request, pk):
    driver = get_object_or_404(Driver, pk=pk)
    return render(request, 'admin/driver_detail.html', {'driver': driver})


# ─── Putyovkalar ──────────────────────────────────────────────────────────────

@admin_required
def admin_putyovkalar(request):
    q = request.GET.get('q', '')
    status = request.GET.get('status', '')
    putyovkalar = Putyovka.objects.select_related('driver').all()
    if q:
        putyovkalar = putyovkalar.filter(
            Q(driver__ism__icontains=q) | Q(raqam__icontains=q) | Q(shifr__icontains=q)
        )
    if status:
        putyovkalar = putyovkalar.filter(status=status)

    paginator = Paginator(putyovkalar, 20)
    page_obj = paginator.get_page(request.GET.get('page'))

    return render(request, 'admin/putyovkalar.html', {
        'putyovkalar': page_obj,
        'search_q': q,
    })


@admin_required
def admin_putyovka_add(request):
    driver_id = request.GET.get('driver_id')
    initial = {}
    if driver_id:
        initial['driver'] = driver_id

    if request.method == 'POST':
        form = PutyovkaForm(request.POST)
        if form.is_valid():
            obj = form.save()
            create_initial_payment(obj, 'putyovka')
            messages.success(request, "Putyovka qo'shildi!")
            return redirect('admin_putyovkalar')
    else:
        form = PutyovkaForm(initial=initial)
    return render(request, 'admin/doc_form.html', {
        'form': form, 'title': "Putyovka qo'shish", 'back_url': 'admin_putyovkalar'
    })


@admin_required
def admin_putyovka_edit(request, pk):
    putyovka = get_object_or_404(Putyovka, pk=pk)
    if request.method == 'POST':
        form = PutyovkaForm(request.POST, instance=putyovka)
        if form.is_valid():
            form.save()
            messages.success(request, "Putyovka yangilandi!")
            return redirect('admin_putyovkalar')
    else:
        form = PutyovkaForm(instance=putyovka)
    return render(request, 'admin/doc_form.html', {
        'form': form, 'title': "Putyovkani tahrirlash", 'back_url': 'admin_putyovkalar'
    })


@admin_required
def admin_putyovka_delete(request, pk):
    putyovka = get_object_or_404(Putyovka, pk=pk)
    putyovka.status = 'tugagan'
    putyovka.save()
    messages.success(request, "Putyovka tugallandi!")
    return redirect('admin_putyovkalar')


@admin_required
def admin_putyovka_tolov(request, pk):
    putyovka = get_object_or_404(Putyovka, pk=pk)
    if request.method == 'POST':
        form = TolovForm(request.POST)
        if form.is_valid():
            summa = form.cleaned_data['summa']
            putyovka.tolangan += summa
            putyovka.save()
            TolovQayd.objects.create(
                driver=putyovka.driver, tur='putyovka',
                summa=summa, izoh=form.cleaned_data.get('izoh', '')
            )
            messages.success(request, f"{summa:,.0f} so'm to'lov qabul qilindi!")
            return redirect('admin_putyovkalar')
    else:
        form = TolovForm()
    return render(request, 'admin/tolov_form.html', {
        'form': form, 'obj': putyovka, 'title': f"Putyovka to'lovi — {putyovka.driver.ism}"
    })


# ─── TIR ──────────────────────────────────────────────────────────────────────

@admin_required
def admin_tirlar(request):
    q = request.GET.get('q', '')
    tirlar = TIR.objects.select_related('driver').all()
    if q:
        tirlar = tirlar.filter(Q(driver__ism__icontains=q) | Q(tir_raqam__icontains=q))

    paginator = Paginator(tirlar, 20)
    page_obj = paginator.get_page(request.GET.get('page'))
    return render(request, 'admin/tirlar.html', {'tirlar': page_obj, 'search_q': q})


@admin_required
def admin_tir_add(request):
    driver_id = request.GET.get('driver_id')
    initial = {'driver': driver_id} if driver_id else {}
    if request.method == 'POST':
        form = TIRForm(request.POST)
        if form.is_valid():
            obj = form.save()
            create_initial_payment(obj, 'tir')
            messages.success(request, "TIR qo'shildi!")
            return redirect('admin_tirlar')
    else:
        form = TIRForm(initial=initial)
    return render(request, 'admin/doc_form.html', {'form': form, 'title': "TIR qo'shish", 'back_url': 'admin_tirlar'})


@admin_required
def admin_tir_edit(request, pk):
    tir = get_object_or_404(TIR, pk=pk)
    if request.method == 'POST':
        form = TIRForm(request.POST, instance=tir)
        if form.is_valid():
            form.save()
            messages.success(request, "TIR yangilandi!")
            return redirect('admin_tirlar')
    else:
        form = TIRForm(instance=tir)
    return render(request, 'admin/doc_form.html', {'form': form, 'title': "TIRni tahrirlash", 'back_url': 'admin_tirlar'})


@admin_required
def admin_tir_delete(request, pk):
    get_object_or_404(TIR, pk=pk).delete()
    messages.success(request, "TIR o'chirildi!")
    return redirect('admin_tirlar')


# ─── Dazvol ───────────────────────────────────────────────────────────────────

@admin_required
def admin_dazvollar(request):
    q = request.GET.get('q', '')
    dazvollar = Dazvol.objects.select_related('driver').all()
    if q:
        dazvollar = dazvollar.filter(Q(driver__ism__icontains=q) | Q(mamlakat__icontains=q))

    paginator = Paginator(dazvollar, 20)
    page_obj = paginator.get_page(request.GET.get('page'))
    return render(request, 'admin/dazvollar.html', {'dazvollar': page_obj, 'search_q': q})


@admin_required
def admin_dazvol_add(request):
    driver_id = request.GET.get('driver_id')
    initial = {'driver': driver_id} if driver_id else {}
    if request.method == 'POST':
        form = DazvolForm(request.POST)
        if form.is_valid():
            obj = form.save()
            create_initial_payment(obj, 'dazvol')
            messages.success(request, "Dazvol qo'shildi!")
            return redirect('admin_dazvollar')
    else:
        form = DazvolForm(initial=initial)
    return render(request, 'admin/doc_form.html', {'form': form, 'title': "Dazvol qo'shish", 'back_url': 'admin_dazvollar'})


@admin_required
def admin_dazvol_edit(request, pk):
    dazvol = get_object_or_404(Dazvol, pk=pk)
    if request.method == 'POST':
        form = DazvolForm(request.POST, instance=dazvol)
        if form.is_valid():
            form.save()
            messages.success(request, "Dazvol yangilandi!")
            return redirect('admin_dazvollar')
    else:
        form = DazvolForm(instance=dazvol)
    return render(request, 'admin/doc_form.html', {'form': form, 'title': "Dazvolni tahrirlash", 'back_url': 'admin_dazvollar'})


@admin_required
def admin_dazvol_delete(request, pk):
    get_object_or_404(Dazvol, pk=pk).delete()
    messages.success(request, "Dazvol o'chirildi!")
    return redirect('admin_dazvollar')


# ─── Litsenziya ───────────────────────────────────────────────────────────────

@admin_required
def admin_litsenziyalar(request):
    q = request.GET.get('q', '')
    litsenziyalar = Litsenziya.objects.select_related('driver').all()
    if q:
        litsenziyalar = litsenziyalar.filter(
            Q(driver__ism__icontains=q) | Q(litsenziya_raqam__icontains=q)
        )
    paginator = Paginator(litsenziyalar, 20)
    page_obj = paginator.get_page(request.GET.get('page'))
    return render(request, 'admin/litsenziyalar.html', {'litsenziyalar': page_obj, 'search_q': q})


@admin_required
def admin_litsenziya_add(request):
    driver_id = request.GET.get('driver_id')
    initial = {'driver': driver_id} if driver_id else {}
    if request.method == 'POST':
        form = LitsenziyaForm(request.POST)
        if form.is_valid():
            obj = form.save()
            create_initial_payment(obj, 'litsenziya')
            messages.success(request, "Litsenziya qo'shildi!")
            return redirect('admin_litsenziyalar')
    else:
        form = LitsenziyaForm(initial=initial)
    return render(request, 'admin/doc_form.html', {'form': form, 'title': "Litsenziya qo'shish", 'back_url': 'admin_litsenziyalar'})


@admin_required
def admin_litsenziya_edit(request, pk):
    litsenziya = get_object_or_404(Litsenziya, pk=pk)
    if request.method == 'POST':
        form = LitsenziyaForm(request.POST, instance=litsenziya)
        if form.is_valid():
            form.save()
            messages.success(request, "Litsenziya yangilandi!")
            return redirect('admin_litsenziyalar')
    else:
        form = LitsenziyaForm(instance=litsenziya)
    return render(request, 'admin/doc_form.html', {'form': form, 'title': "Litsenziyani tahrirlash", 'back_url': 'admin_litsenziyalar'})


@admin_required
def admin_litsenziya_delete(request, pk):
    get_object_or_404(Litsenziya, pk=pk).delete()
    messages.success(request, "Litsenziya o'chirildi!")
    return redirect('admin_litsenziyalar')


# ─── Ijara ────────────────────────────────────────────────────────────────────

@admin_required
def admin_ijara(request):
    q = request.GET.get('q', '')
    ijaralar = IjaraShartnoma.objects.select_related('driver').all()
    if q:
        ijaralar = ijaralar.filter(Q(driver__ism__icontains=q) | Q(manzil__icontains=q))
    paginator = Paginator(ijaralar, 20)
    page_obj = paginator.get_page(request.GET.get('page'))
    return render(request, 'admin/ijara.html', {'ijaralar': page_obj, 'search_q': q})


@admin_required
def admin_ijara_add(request):
    driver_id = request.GET.get('driver_id')
    initial = {'driver': driver_id} if driver_id else {}
    if request.method == 'POST':
        form = IjaraForm(request.POST)
        if form.is_valid():
            obj = form.save()
            create_initial_payment(obj, 'ijara')
            messages.success(request, "Ijara shartnomasi qo'shildi!")
            return redirect('admin_ijara')
    else:
        form = IjaraForm(initial=initial)
    return render(request, 'admin/doc_form.html', {'form': form, 'title': "Ijara shartnomasi qo'shish", 'back_url': 'admin_ijara'})


@admin_required
def admin_ijara_edit(request, pk):
    ijara = get_object_or_404(IjaraShartnoma, pk=pk)
    if request.method == 'POST':
        form = IjaraForm(request.POST, instance=ijara)
        if form.is_valid():
            form.save()
            messages.success(request, "Ijara yangilandi!")
            return redirect('admin_ijara')
    else:
        form = IjaraForm(instance=ijara)
    return render(request, 'admin/doc_form.html', {'form': form, 'title': "Ijarani tahrirlash", 'back_url': 'admin_ijara'})


@admin_required
def admin_ijara_delete(request, pk):
    get_object_or_404(IjaraShartnoma, pk=pk).delete()
    messages.success(request, "Ijara shartnomasi o'chirildi!")
    return redirect('admin_ijara')


# ─── Chiqimlar ────────────────────────────────────────────────────────────────

@admin_required
def admin_chiqimlar(request):
    q = request.GET.get('q', '')
    chiqimlar = Chiqim.objects.select_related('driver').all()
    if q:
        chiqimlar = chiqimlar.filter(
            Q(driver__ism__icontains=q) | Q(izoh__icontains=q)
        )
    jami = chiqimlar.aggregate(s=Sum('summa'))['s'] or 0
    paginator = Paginator(chiqimlar, 20)
    page_obj = paginator.get_page(request.GET.get('page'))
    return render(request, 'admin/chiqimlar.html', {
        'chiqimlar': page_obj, 'search_q': q, 'jami': jami
    })


@admin_required
def admin_chiqim_add(request):
    if request.method == 'POST':
        form = ChiqimForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Chiqim qo'shildi!")
            return redirect('admin_chiqimlar')
    else:
        form = ChiqimForm()
    return render(request, 'admin/doc_form.html', {'form': form, 'title': "Chiqim qo'shish", 'back_url': 'admin_chiqimlar'})


@admin_required
def admin_chiqim_delete(request, pk):
    get_object_or_404(Chiqim, pk=pk).delete()
    messages.success(request, "Chiqim o'chirildi!")
    return redirect('admin_chiqimlar')


# ─── Word Export ──────────────────────────────────────────────────────────────

# Bu funksiyalarni core/views.py ichidagi mavjudlarini almashtiring
# (taxminan 690-760 qatorlar oralig'ida)

import os


def _replace_in_paragraph_xml(para, replacements):
    """
    Paragraph ichidagi matnni almashtiradi, lekin run-lardagi
    formatlashni (shrift, bold, italic) saqlab qoladi.
    """
    if not para.runs:
        return

    # 1-bosqich: har bir run ichida alohida almashtirish (formatlash saqlanadi)
    for run in para.runs:
        for old, new in replacements.items():
            if old and old in run.text:
                run.text = run.text.replace(old, new)

    # 2-bosqich: placeholder bir nechta run-ga bo'lingan bo'lsa
    full_text = ''.join(r.text for r in para.runs)
    remaining = [old for old in replacements if old and old in full_text]
    if not remaining:
        return

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    # Faqat kerakli ketma-ket run-larni birlashtiramiz
    for old in remaining:
        new = replacements[old]
        runs = para.runs
        texts = [r.text for r in runs]
        joined = ''.join(texts)
        idx = joined.find(old)
        if idx < 0:
            continue

        # qaysi run-lar shu oraliqqa tushishini topamiz
        pos = 0
        start_run = end_run = None
        start_off = end_off = 0
        for i, t in enumerate(texts):
            nxt = pos + len(t)
            if start_run is None and idx < nxt:
                start_run = i
                start_off = idx - pos
            if idx + len(old) <= nxt:
                end_run = i
                end_off = idx + len(old) - pos
                break
            pos = nxt
        if start_run is None or end_run is None:
            continue

        # birinchi run-da matnni almashtiramiz, formatlashni saqlaymiz
        first = runs[start_run]
        last = runs[end_run]
        first.text = texts[start_run][:start_off] + new
        # oxirgi run-dagi qoldiqni keyinroq qo'shamiz
        tail = texts[end_run][end_off:]
        # oradagi run-larni o'chiramiz
        for j in range(end_run, start_run, -1):
            r_el = runs[j]._element
            r_el.getparent().remove(r_el)
        # tail bo'lsa, yangi run sifatida qo'shamiz (oxirgi run formatlashi bilan)
        if tail:
            new_r = copy.deepcopy(last._element)
            # ichidagi <w:t>ni tail bilan almashtiramiz
            for t_el in new_r.findall(qn('w:t')):
                new_r.remove(t_el)
            t_new = OxmlElement('w:t')
            t_new.text = tail
            t_new.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_r.append(t_new)
            first._element.addnext(new_r)


def _replace_in_doc(doc, replacements):
    """Hujjatdagi paragraph va jadval cell-laridagi matnni almashtiradi."""
    for para in doc.paragraphs:
        _replace_in_paragraph_xml(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph_xml(para, replacements)


def _replace_name_in_doc(doc, old_parts, new_name):
    """
    Ism almashtirish: avval to'liq ismni, keyin alohida qismlarini.
    Formatlash (bold, shrift) saqlab qolinadi.
    Shablondagi probel-asosli ustun layoutini buzmaslik uchun
    yangi ism eski ism uzunligigacha probel bilan to'ldiriladi.
    """
    full_old = ' '.join(old_parts)
    # Layoutni saqlash: agar yangi ism qisqa bo'lsa, probel bilan to'ldiramiz
    if len(new_name) < len(full_old):
        padded = new_name + ' ' * (len(full_old) - len(new_name))
    else:
        padded = new_name
    _replace_in_doc(doc, {full_old: padded})

    new_parts = new_name.split()
    for i, part in enumerate(old_parts):
        replacement = new_parts[i] if i < len(new_parts) else ''
        if part and part != replacement:
            _replace_in_doc(doc, {part: replacement})


@admin_required
def putyovka_word(request, pk):
    putyovka = get_object_or_404(Putyovka, pk=pk)
    try:
        from docx import Document
        import io

        template_path = os.path.join(os.path.dirname(__file__), 'putyovka.docx')
        doc = Document(template_path)

        replacements = {
            '${start}': putyovka.boshlanish.strftime('%d.%m.%Y'),
            '${end}': putyovka.tugash.strftime('%d.%m.%Y'),
            '${car_name}': putyovka.driver.mashina or '',
            '${car_number}': putyovka.driver.raqam or '',
            '${driver}': putyovka.driver.ism or '',
        }
        # Pritsep raqami (agar shablonda alohida placeholder bo'lsa)
        if putyovka.raqam:
            replacements['${pritsep}'] = putyovka.raqam
        _replace_in_doc(doc, replacements)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        safe_name = putyovka.driver.ism.replace(' ', '_')
        response['Content-Disposition'] = f'attachment; filename="putyovka_{safe_name}_{putyovka.id}.docx"'
        return response
    except Exception as e:
        messages.error(request, f"Xatolik: {e}")
        return redirect('admin_driver_detail', pk=putyovka.driver.pk)


@admin_required
def driver_blanka_word(request, pk):
    driver = get_object_or_404(Driver, pk=pk)
    try:
        from docx import Document
        import io

        template_path = os.path.join(os.path.dirname(__file__), 'blanka.docx')
        doc = Document(template_path)

        old_parts = ['Ismonov', 'Doniyorbek', 'Shavkatovich']
        _replace_name_in_doc(doc, old_parts, driver.ism)

        # Pasport seriya almashtirish
        new_pasport = driver.pasport_seriya.strip() if driver.pasport_seriya else 'FA 1339223'
        _replace_in_doc(doc, {'FA 1339223': new_pasport})

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        safe_name = driver.ism.replace(' ', '_')
        response['Content-Disposition'] = f'attachment; filename="blanka_{safe_name}.docx"'
        return response
    except Exception as e:
        messages.error(request, f"Xatolik: {e}")
        return redirect('admin_driver_detail', pk=pk)


@admin_required
def driver_doverennost_word(request, pk):
    driver = get_object_or_404(Driver, pk=pk)
    try:
        from docx import Document
        import io

        template_path = os.path.join(os.path.dirname(__file__), 'doverennost.docx')
        doc = Document(template_path)

        old_parts = ['Ismonov', 'Doniyorbek', 'Shavkatovich']
        _replace_name_in_doc(doc, old_parts, driver.ism)

        # Pasport seriya: GET parametr > driver.pasport_seriya > default
        pasport_param = (request.GET.get('pasport') or '').strip()
        new_pasport = pasport_param or (driver.pasport_seriya.strip() if driver.pasport_seriya else 'FA 1339223')
        _replace_in_doc(doc, {'FA 1339223': new_pasport})

        # Boshlanish va tugash sanalari — GET parametri yoki default (bugun, +3 yil)
        from datetime import date as _d, timedelta as _td, datetime as _dt
        def _parse(s):
            try:
                return _dt.strptime(s, '%Y-%m-%d').date()
            except Exception:
                return None
        start = _parse(request.GET.get('start', '')) or _d.today()
        end = _parse(request.GET.get('end', ''))
        if not end:
            try:
                end = start.replace(year=start.year + 3) - _td(days=1)
            except ValueError:
                end = start + _td(days=365 * 3 - 1)

        date_replacements = {
            '«20» 04. 2026': f'«{start.day:02d}» {start.month:02d}. {start.year}',
            '«19»04. 2029':  f'«{end.day:02d}»{end.month:02d}. {end.year}',
            '20» 04. 2026': f'{start.day:02d}» {start.month:02d}. {start.year}',
            '19»04. 2029':  f'{end.day:02d}»{end.month:02d}. {end.year}',
        }
        _replace_in_doc(doc, date_replacements)


        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        safe_name = driver.ism.replace(' ', '_')
        response['Content-Disposition'] = f'attachment; filename="doverennost_{safe_name}.docx"'
        return response
    except Exception as e:
        messages.error(request, f"Xatolik: {e}")
        return redirect('admin_driver_detail', pk=pk)


# ─── Driver Views ─────────────────────────────────────────────────────────────

@driver_required
def driver_dashboard(request):
    driver_id = request.session.get('driver_id')
    driver = get_object_or_404(Driver, pk=driver_id)
    today = date.today()
    soon_7 = today + timedelta(days=7)

    # Muddati 7 kun ichida tugaydigan hujjatlar
    expiring_soon = []
    p = driver.aktiv_putyovka
    if p and not p.is_expired and p.tugash <= soon_7:
        expiring_soon.append({"tur": "Putyovka", "kun": (p.tugash - today).days, "sana": p.tugash.strftime("%d.%m.%Y")})
    t = driver.aktiv_tir
    if t and not t.is_expired and t.tugash <= soon_7:
        expiring_soon.append({"tur": "TIR", "kun": (t.tugash - today).days, "sana": t.tugash.strftime("%d.%m.%Y")})
    dz = driver.aktiv_dazvol
    if dz and not dz.is_expired and dz.tugash <= soon_7:
        expiring_soon.append({"tur": "Dazvol", "kun": (dz.tugash - today).days, "sana": dz.tugash.strftime("%d.%m.%Y")})

    # Muddati otgan hujjatlar
    expired_docs = []
    if p and p.is_expired:
        expired_docs.append({"tur": "Putyovka", "sana": p.tugash.strftime("%d.%m.%Y")})
    if t and t.is_expired:
        expired_docs.append({"tur": "TIR", "sana": t.tugash.strftime("%d.%m.%Y")})
    if dz and dz.is_expired:
        expired_docs.append({"tur": "Dazvol", "sana": dz.tugash.strftime("%d.%m.%Y")})

    admin_phone = "+998 90 000 00 00"

    return render(request, "driver/dashboard.html", {
        "driver": driver,
        "expiring_soon": expiring_soon,
        "expired_docs": expired_docs,
        "admin_phone": admin_phone,
    })

# ─── Admin management ─────────────────────────────────────────────────────────

@admin_required
def admin_add_admin(request):
    """Yangi admin qo'shish va adminlar ro'yxati"""
    error = None
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '').strip()
        email    = request.POST.get('email', '').strip()
        if not username or not password:
            error = "Login va parol majburiy!"
        elif User.objects.filter(username=username).exists():
            error = f"'{username}' login allaqachon mavjud!"
        else:
            User.objects.create_superuser(username=username, password=password, email=email)
            messages.success(request, f"Admin '{username}' muvaffaqiyatli qo'shildi!")
            return redirect('admin_add_admin')
    admins = User.objects.filter(is_staff=True).order_by('-date_joined')
    return render(request, 'admin/add_admin.html', {'error': error, 'admins': admins})


@admin_required
def admin_driver_sms(request, pk):
    """Haydovchiga faol hujjatlar haqida SMS yuborish"""
    import json
    import urllib.request
    import urllib.error
    from datetime import date

    driver = get_object_or_404(Driver, pk=pk)

    if request.method != 'POST':
        return redirect('admin_driver_detail', pk=pk)

    # Faol hujjatlarni yig'amiz
    lines = []
    today = date.today()

    # Putyovka
    for p in driver.putyovkalar.exclude(status='tugagan'):
        qarz_str = f"{int(p.qarz):,} so'm qarz".replace(',', ' ') if p.qarz > 0 else "qarz yo'q"
        lines.append(
            f"📄 Putyovka: {p.boshlanish.strftime('%d.%m.%Y')} - {p.tugash.strftime('%d.%m.%Y')}, {qarz_str}"
        )

    # TIR
    for t in driver.tirlar.exclude(status='tugagan'):
        qarz_str = f"{int(t.qarz):,} so'm qarz".replace(',', ' ') if t.qarz > 0 else "qarz yo'q"
        lines.append(
            f"🚛 TIR ({t.tir_raqam}): {t.boshlanish.strftime('%d.%m.%Y')} - {t.tugash.strftime('%d.%m.%Y')}, {qarz_str}"
        )

    # Dazvol
    for dz in driver.dazvollar.exclude(status='tugagan'):
        qarz_str = f"{int(dz.qarz):,} so'm qarz".replace(',', ' ') if dz.qarz > 0 else "qarz yo'q"
        lines.append(
            f"📝 Dazvol ({dz.mamlakat}): {dz.boshlanish.strftime('%d.%m.%Y')} - {dz.tugash.strftime('%d.%m.%Y')}, {qarz_str}"
        )

    # Litsenziya
    for l in driver.litsenziyalar.exclude(status='tugagan'):
        qarz_str = f"{int(l.qarz):,} so'm qarz".replace(',', ' ') if l.qarz > 0 else "qarz yo'q"
        lines.append(
            f"🛡 Litsenziya: {l.boshlanish.strftime('%d.%m.%Y')} - {l.tugash.strftime('%d.%m.%Y')}, {qarz_str}"
        )

    # Ijara
    for ij in driver.ijara_shartnomalari.exclude(status='tugagan'):
        qarz_str = f"{int(ij.qarz):,} so'm qarz".replace(',', ' ') if ij.qarz > 0 else "qarz yo'q"
        lines.append(
            f"🏠 Ijara ({ij.manzil}): {ij.boshlanish.strftime('%d.%m.%Y')} - {ij.tugash.strftime('%d.%m.%Y')}, {qarz_str}"
        )

    if not lines:
        messages.warning(request, "Haydovchida faol hujjat mavjud emas.")
        return redirect('admin_driver_detail', pk=pk)

    jami_qarz = driver.jami_qarz
    jami_str = f"{int(jami_qarz):,} so'm".replace(',', ' ')

    sms_text = f"Hurmatli {driver.ism},\nFaol hujjatlaringiz:\n"
    sms_text += "\n".join(lines)
    sms_text += f"\n\nJami qarz: {jami_str}"

    # Telefon raqamini tozalash: +998901234567 -> 998901234567
    telefon = driver.telefon.strip().replace(' ', '').replace('-', '')
    if telefon.startswith('+'):
        telefon = telefon[1:]

    # ── Eskiz.uz SMS Shlyuz API ───────────────────────────────────────────────
    from django.conf import settings
    eskiz_email    = getattr(settings, 'ESKIZ_EMAIL', '')
    eskiz_password = getattr(settings, 'ESKIZ_PASSWORD', '')

    if not eskiz_email or not eskiz_password:
        messages.error(request, "SMS xizmati sozlanmagan. settings.py ga ESKIZ_EMAIL va ESKIZ_PASSWORD qo'shing.")
        return redirect('admin_driver_detail', pk=pk)

    try:
        # 1. Token olish
        auth_data = json.dumps({
            "email": eskiz_email,
            "password": eskiz_password
        }).encode('utf-8')
        auth_req = urllib.request.Request(
            "https://notify.eskiz.uz/api/auth/login",
            data=auth_data,
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        with urllib.request.urlopen(auth_req, timeout=10) as resp:
            auth_resp = json.loads(resp.read().decode('utf-8'))
        token = auth_resp.get('data', {}).get('token', '')

        if not token:
            err = auth_resp.get('message', str(auth_resp))
            messages.error(request, f"Token olishda xatolik: {err}")
            return redirect('admin_driver_detail', pk=pk)

        # 2. SMS yuborish — telefon: 998XXXXXXXXX formatida
        # Eskiz test rejimi: faqat "Bu Eskiz dan test" matni ishlaydi
        # Akkount to'ldirilgandan so'ng: yuboriluvchi_matn = sms_text
        yuboriluvchi_matn = "Bu Eskiz dan test"
        sms_data = json.dumps({
            "mobile_phone": telefon,
            "message": yuboriluvchi_matn,
            "from": "4546",
            "callback_url": ""
        }).encode('utf-8')
        sms_req = urllib.request.Request(
            "https://notify.eskiz.uz/api/message/sms/send",
            data=sms_data,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {token}"
            },
            method="POST"
        )
        with urllib.request.urlopen(sms_req, timeout=10) as resp:
            sms_resp = json.loads(resp.read().decode('utf-8'))

        status = sms_resp.get('status', '')
        if status in ('waiting', 'sent', 'ok'):
            messages.success(request, f"SMS yuborildi: {driver.telefon}")
        else:
            err_msg = sms_resp.get('message', str(sms_resp))
            messages.error(request, f"SMS xatolik: {err_msg} | Javob: {sms_resp}")

    except urllib.error.HTTPError as e:
        body = e.read().decode('utf-8', errors='replace')
        messages.error(request, f"HTTP {e.code} xatolik: {body}")
    except urllib.error.URLError as e:
        messages.error(request, f"Ulanish xatoligi: {e.reason}")
    except Exception as e:
        messages.error(request, f"Kutilmagan xatolik: {str(e)}")

    return redirect('admin_driver_detail', pk=pk)


@admin_required
def admin_delete_admin(request, pk):
    if request.user.pk == pk:
        messages.error(request, "O'zingizni o'chira olmaysiz!")
        return redirect('admin_add_admin')
    user = get_object_or_404(User, pk=pk, is_staff=True)
    user.delete()
    messages.success(request, "Admin o'chirildi!")
    return redirect('admin_add_admin')


@admin_required
def admin_tir_tolov(request, pk):
    obj=get_object_or_404(TIR, pk=pk)
    if request.method=="POST":
        form=TolovForm(request.POST)
        if form.is_valid():
            summa=form.cleaned_data['summa']
            obj.tolangan+=summa
            obj.save()
            TolovQayd.objects.create(driver=obj.driver,tur='tir',summa=summa,izoh=form.cleaned_data.get('izoh',''))
            messages.success(request,"To'lov qabul qilindi!")
            return redirect('admin_tirlar')
    else: form=TolovForm()
    return render(request,'admin/tolov_form.html',{'form':form,'obj':obj,'title':f"TIR to'lovi — {obj.driver.ism}"})

@admin_required
def admin_dazvol_tolov(request, pk):
    obj=get_object_or_404(Dazvol, pk=pk)
    if request.method=="POST":
        form=TolovForm(request.POST)
        if form.is_valid():
            summa=form.cleaned_data['summa']; obj.tolangan+=summa; obj.save()
            TolovQayd.objects.create(driver=obj.driver,tur='dazvol',summa=summa,izoh=form.cleaned_data.get('izoh',''))
            return redirect('admin_dazvollar')
    else: form=TolovForm()
    return render(request,'admin/tolov_form.html',{'form':form,'obj':obj,'title':f"Dazvol to'lovi — {obj.driver.ism}"})

@admin_required
def admin_litsenziya_tolov(request, pk):
    obj=get_object_or_404(Litsenziya, pk=pk)
    if request.method=="POST":
        form=TolovForm(request.POST)
        if form.is_valid():
            summa=form.cleaned_data['summa']; obj.tolangan+=summa; obj.save()
            TolovQayd.objects.create(driver=obj.driver,tur='litsenziya',summa=summa,izoh=form.cleaned_data.get('izoh',''))
            return redirect('admin_litsenziyalar')
    else: form=TolovForm()
    return render(request,'admin/tolov_form.html',{'form':form,'obj':obj,'title':f"Litsenziya to'lovi — {obj.driver.ism}"})

@admin_required
def admin_ijara_tolov(request, pk):
    obj=get_object_or_404(IjaraShartnoma, pk=pk)
    if request.method=="POST":
        form=TolovForm(request.POST)
        if form.is_valid():
            summa=form.cleaned_data['summa']; obj.tolangan+=summa; obj.save()
            TolovQayd.objects.create(driver=obj.driver,tur='ijara',summa=summa,izoh=form.cleaned_data.get('izoh',''))
            return redirect('admin_ijara')
    else: form=TolovForm()
    return render(request,'admin/tolov_form.html',{'form':form,'obj':obj,'title':f"Ijara to'lovi — {obj.driver.ism}"})
